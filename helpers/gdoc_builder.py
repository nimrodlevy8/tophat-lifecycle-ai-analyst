"""Google Doc builder — generates a .docx Analysis Readout from structured data.

Uses python-docx to create a richly formatted Word document that can be
uploaded to Google Drive and converted to a Google Doc.

Template: Analysis Readout
  H1 Title → H2 Context → H2 Summary (Primary Learnings + Recommendations)
  → H2 Analysis (findings with charts) → H2 Next Steps → H2 Resources (SQL)
"""

from __future__ import annotations

import logging
import os
import traceback
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Debug logger — writes to working/export_gdoc_debug.log
_log = logging.getLogger("gdoc_builder")


def _setup_debug_log(working_dir: str = "working") -> None:
    """Configure file logging for debug output."""
    if _log.handlers:
        return
    os.makedirs(working_dir, exist_ok=True)
    handler = logging.FileHandler(
        os.path.join(working_dir, "export_gdoc_debug.log"), mode="a"
    )
    handler.setFormatter(logging.Formatter(
        "%(asctime)s [%(levelname)s] %(message)s"
    ))
    _log.addHandler(handler)
    _log.setLevel(logging.DEBUG)


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class SubFinding:
    """A sub-finding within a finding (renders as H4)."""
    title: str
    body: str
    chart_path: Optional[str] = None  # absolute path to PNG
    chart_caption: Optional[str] = None


@dataclass
class Finding:
    """A top-level finding (renders as H3 under Analysis)."""
    headline: str  # action headline
    summary: str
    sub_findings: list[SubFinding] = field(default_factory=list)
    data_stamp: Optional[str] = None        # "[1.5M rows | Jan-Mar 2026 | ORDERS | A (93)]"
    methodology: Optional[str] = None       # "Approach: variance decomposition..."
    sql: Optional[str] = None               # full SQL query text
    cross_verification: Optional[str] = None # "Type B: Parts-to-whole -- PASS"


@dataclass
class Recommendation:
    """A recommendation with confidence level."""
    action: str  # "Action verb + scope"
    rationale: str
    confidence: str = "Medium"  # High / Medium / Low


@dataclass
class SqlQuery:
    """A SQL query for the Resources section."""
    title: str
    sql: str
    used_in_finding: Optional[int] = None  # 1-indexed finding number
    database: Optional[str] = None


@dataclass
class SuccessTracking:
    """Success metric for Next Steps."""
    metric: str
    baseline: str
    target: str
    check_in_date: Optional[str] = None


@dataclass
class AnalysisData:
    """All data needed to build an Analysis Readout document."""
    title: str
    subtitle: Optional[str] = None
    author: str = "AI Analyst"
    date: str = ""
    confidence_grade: Optional[str] = None  # e.g., "B+"
    confidence_score: Optional[int] = None  # e.g., 85
    confidence_caveat: Optional[str] = None
    context: Optional[str] = None
    findings: list[Finding] = field(default_factory=list)
    synthesis: Optional[str] = None
    implications: Optional[str] = None
    recommendations: list[Recommendation] = field(default_factory=list)
    next_steps_actions: Optional[str] = None
    success_tracking: Optional[SuccessTracking] = None
    open_questions: Optional[str] = None
    sql_queries: list[SqlQuery] = field(default_factory=list)
    companion_analyses: Optional[str] = None
    data_sources: Optional[str] = None
    provenance_blocks: Optional[list] = None  # from provenance_assembler


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_BODY_FONT = "Calibri"
_BODY_SIZE = Pt(11)
_HEADING_COLOR = RGBColor(0x1E, 0x29, 0x3B)  # dark navy
_BODY_COLOR = RGBColor(0x22, 0x22, 0x22)
_CAPTION_COLOR = RGBColor(0x66, 0x66, 0x66)
_SQL_BG_COLOR = "F2F2F2"
_SQL_FONT = "Courier New"
_SQL_SIZE = Pt(9)
_CHART_WIDTH = Inches(6.0)

# Bold label prefixes to detect and style
BOLD_LABELS = [
    "The Insight:",
    "Why this matters for product:",
    "Bottom line:",
    "Key context:",
    "Data quality flag:",
    "Sample size warning:",
    "The creative angle:",
    "The interpretation:",
]


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex: str) -> None:
    """Set background shading on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_bookmark(paragraph, bookmark_name: str) -> None:
    """Add a bookmark anchor to a paragraph."""
    # Unique ID for the bookmark pair
    bm_id = str(abs(hash(bookmark_name)) % 10**8)

    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), bm_id)
    bm_start.set(qn("w:name"), bookmark_name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), bm_id)

    paragraph._p.append(bm_start)
    paragraph._p.append(bm_end)


def _add_hyperlink(paragraph, text: str, bookmark_name: str) -> None:
    """Add an internal hyperlink (to a bookmark) inside a paragraph."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)

    run_elem = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    # Blue underline style for links
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run_elem.append(rpr)

    text_elem = OxmlElement("w:t")
    text_elem.set(qn("xml:space"), "preserve")
    text_elem.text = text
    run_elem.append(text_elem)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _style_run(run, font_name=_BODY_FONT, size=_BODY_SIZE,
               color=_BODY_COLOR, bold=False, italic=False) -> None:
    """Apply consistent styling to a run."""
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def _add_body_paragraph(doc, text: str, bold=False, italic=False,
                        alignment=None) -> None:
    """Add a styled body paragraph, rendering inline **bold** markdown
    and auto-bolding known label prefixes."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment

    # Check if text starts with a known bold label
    label_found = None
    if not bold:
        for label in BOLD_LABELS:
            if text.startswith(label):
                label_found = label
                break

    if label_found:
        run_label = p.add_run(label_found)
        _style_run(run_label, bold=True)
        remaining = text[len(label_found):]
        # Render remaining with inline **bold** support
        _add_markdown_runs(p, remaining, italic=italic)
    elif bold:
        # Whole paragraph bold — still handle inline markdown
        _add_markdown_runs(p, text, base_bold=True, italic=italic)
    else:
        # Render with inline **bold** support
        _add_markdown_runs(p, text, italic=italic)

    return p


def _add_markdown_runs(paragraph, text: str, base_bold=False,
                        italic=False) -> None:
    """Split text on **bold** markers and add runs with correct styling.

    Handles `**text**` inline bold syntax from markdown.
    """
    import re
    # Split on **...**  — parts alternate: plain, bold, plain, bold, ...
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        is_bold_segment = (idx % 2 == 1)  # odd indices are inside **...**
        run = paragraph.add_run(part)
        _style_run(run, bold=(base_bold or is_bold_segment), italic=italic)


def _add_heading(doc, text: str, level: int, bookmark_name: str = None):
    """Add a heading with optional bookmark anchor."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _HEADING_COLOR
    if bookmark_name:
        _add_bookmark(h, bookmark_name)
    return h


def _add_spacing_paragraph(doc) -> None:
    """Add an empty paragraph for spacing."""
    doc.add_paragraph()


def _confidence_sort_key(rec: Recommendation) -> int:
    """Sort recommendations: High first, then Medium, then Low."""
    order = {"high": 0, "medium": 1, "low": 2}
    return order.get(rec.confidence.lower(), 1)


# ---------------------------------------------------------------------------
# Provenance helpers (data stamps, methodology, confidence badges)
# ---------------------------------------------------------------------------

# Badge shading colors by grade
_BADGE_COLORS = {
    "A": "D4EDDA",   # green
    "B": "FFF3CD",   # yellow
    "C": "FFE0B2",   # orange
}
_BADGE_DEFAULT_COLOR = "E8E8E8"  # gray


def _add_data_stamp(doc, stamp_text: str) -> None:
    """Add a gray-shaded data stamp cell (like Notion's gray callout).

    Renders a single-row table with gray background, 9pt italic text.
    Skips if stamp_text is empty or starts with "[0 rows".
    """
    if not stamp_text or stamp_text.startswith("[0 rows"):
        return

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_shading(cell, _SQL_BG_COLOR)

    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(stamp_text)
    _style_run(run, size=Pt(9), color=_CAPTION_COLOR, italic=True)


def _add_methodology_sql_section(doc, methodology: str = None,
                                  sql: str = None,
                                  cross_verification: str = None) -> None:
    """Add a Methodology & SQL section under a finding.

    Renders methodology as italic text, SQL in a gray monospace cell,
    and cross-verification as labeled gray italic text.
    Only renders if at least one field is non-None.
    """
    if not any([methodology, sql, cross_verification]):
        return

    h = doc.add_heading("Methodology & SQL", level=4)
    for run in h.runs:
        run.font.color.rgb = _CAPTION_COLOR

    if methodology:
        p = doc.add_paragraph()
        run = p.add_run(methodology)
        _style_run(run, size=Pt(10), color=_BODY_COLOR, italic=True)

    if sql:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        _set_cell_shading(cell, _SQL_BG_COLOR)
        cell.text = ""
        p_sql = cell.paragraphs[0]
        run_sql = p_sql.add_run(sql.strip())
        _style_run(run_sql, font_name=_SQL_FONT, size=_SQL_SIZE)

    if cross_verification:
        p = doc.add_paragraph()
        run_label = p.add_run("Cross-verification: ")
        _style_run(run_label, size=Pt(9), color=_CAPTION_COLOR, bold=True,
                    italic=True)
        run_val = p.add_run(cross_verification)
        _style_run(run_val, size=Pt(9), color=_CAPTION_COLOR, italic=True)


def _add_confidence_badge(doc, grade: str, score: int = None,
                           caveat: str = None) -> None:
    """Add a colored confidence badge as a shaded table cell.

    Green for A, yellow for B, orange for C, gray default.
    """
    badge_text = f"Analysis Confidence: {grade}"
    if score is not None:
        badge_text += f" ({score}/100)"

    color = _BADGE_COLORS.get(grade[0] if grade else "", _BADGE_DEFAULT_COLOR)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_shading(cell, color)

    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(badge_text)
    _style_run(run, bold=True, size=Pt(11))

    if caveat:
        p_caveat = doc.add_paragraph()
        run_caveat = p_caveat.add_run(caveat)
        _style_run(run_caveat, italic=True, size=Pt(10), color=_CAPTION_COLOR)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_title_block(doc: Document, data: AnalysisData) -> None:
    """H1 title, subtitle, author, date, confidence badge."""
    _add_heading(doc, data.title, level=1)

    if data.subtitle:
        _add_body_paragraph(doc, data.subtitle, italic=True)

    if data.author or data.date:
        parts = []
        if data.author:
            parts.append(data.author)
        if data.date:
            parts.append(data.date)
        _add_body_paragraph(doc, " | ".join(parts))

    if data.confidence_grade is not None:
        _add_confidence_badge(doc, data.confidence_grade,
                               data.confidence_score, data.confidence_caveat)

    _add_spacing_paragraph(doc)


def _build_context(doc: Document, data: AnalysisData) -> None:
    """H2 Context section."""
    if not data.context:
        return
    _add_heading(doc, "Context", level=2)
    _add_body_paragraph(doc, data.context)


def _build_summary(doc: Document, data: AnalysisData,
                   figure_counter: list) -> None:
    """H2 Summary with Primary Learnings (>> linked) and Recommendations."""
    _add_heading(doc, "Summary", level=2)

    # Primary Learnings
    if data.findings:
        _add_heading(doc, "Primary Learnings", level=3)

        for i, finding in enumerate(data.findings, 1):
            p = doc.add_paragraph()
            p.style = "List Number"

            # Bold headline
            run_num = p.add_run(f"{finding.headline}")
            _style_run(run_num, bold=True)

            # Detail from summary (first sentence)
            summary_text = finding.summary
            first_sentence = summary_text.split(". ")[0] + "." if ". " in summary_text else summary_text
            if first_sentence and not first_sentence.startswith(finding.headline):
                run_detail = p.add_run(f" \u2014 {first_sentence} ")
                _style_run(run_detail)

            # >> link to finding
            bookmark_name = f"finding-{i}"
            _add_hyperlink(p, f">> See Finding {i}", bookmark_name)

    # Recommendations
    if data.recommendations:
        _add_heading(doc, "Recommendations", level=3)

        sorted_recs = sorted(data.recommendations, key=_confidence_sort_key)
        for i, rec in enumerate(sorted_recs, 1):
            p = doc.add_paragraph()
            p.style = "List Number"

            run_action = p.add_run(rec.action)
            _style_run(run_action, bold=True)

            run_detail = p.add_run(
                f" \u2014 {rec.rationale} Confidence: {rec.confidence}."
            )
            _style_run(run_detail)

    _add_spacing_paragraph(doc)


def _build_analysis(doc: Document, data: AnalysisData,
                    figure_counter: list) -> None:
    """H2 Analysis with findings, sub-findings, charts, synthesis, implications."""
    if not data.findings and not data.synthesis:
        return

    doc.add_page_break()
    _add_heading(doc, "Analysis", level=2)

    for i, finding in enumerate(data.findings, 1):
        bookmark_name = f"finding-{i}"
        _add_heading(doc, f"Finding {i}: {finding.headline}", level=3,
                     bookmark_name=bookmark_name)

        # Data stamp (gray shaded cell)
        if finding.data_stamp:
            _add_data_stamp(doc, finding.data_stamp)

        if finding.summary:
            _add_body_paragraph(doc, finding.summary)

        for sub in finding.sub_findings:
            _add_heading(doc, sub.title, level=4)

            if sub.body:
                # Split body into paragraphs and handle bold labels
                for para_text in sub.body.split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        _add_body_paragraph(doc, para_text)

            # Chart
            if sub.chart_path and os.path.isfile(sub.chart_path):
                try:
                    doc.add_picture(sub.chart_path, width=_CHART_WIDTH)
                    # Center the image
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Caption
                    figure_counter[0] += 1
                    caption = sub.chart_caption or sub.title
                    caption_text = f"Figure {figure_counter[0]}: {caption}."
                    p_cap = _add_body_paragraph(doc, caption_text, italic=True,
                                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    for run in p_cap.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = _CAPTION_COLOR
                except Exception as e:
                    filename = os.path.basename(sub.chart_path)
                    _log.error("Failed to embed chart %s: %s", filename, e)
                    _add_body_paragraph(
                        doc,
                        f"[Chart: {filename} \u2014 could not embed: {e}]",
                        italic=True,
                    )
            elif sub.chart_path:
                # Missing chart placeholder
                filename = os.path.basename(sub.chart_path)
                _log.warning("Chart file not found: %s", sub.chart_path)
                _add_body_paragraph(
                    doc,
                    f"[Chart: {filename} \u2014 file not found]",
                    italic=True,
                )

        # Methodology & SQL section (after all sub-findings)
        _add_methodology_sql_section(
            doc,
            methodology=finding.methodology,
            sql=finding.sql,
            cross_verification=finding.cross_verification,
        )

        _add_spacing_paragraph(doc)

    # Synthesis
    if data.synthesis:
        _add_heading(doc, "Synthesis", level=3)
        _add_body_paragraph(doc, data.synthesis)

        if data.confidence_grade:
            conf_note = f"This analysis carries {data.confidence_grade} confidence"
            if data.confidence_score is not None:
                conf_note += f" ({data.confidence_score}/100)"
            conf_note += "."
            if data.confidence_caveat:
                conf_note += f" {data.confidence_caveat}"
            _add_body_paragraph(doc, conf_note, italic=True)

    # Implications
    if data.implications:
        _add_heading(doc, "Implications", level=3)
        _add_body_paragraph(doc, data.implications)


def _build_next_steps(doc: Document, data: AnalysisData) -> None:
    """H2 Next Steps with actions, success tracking, open questions."""
    has_content = (data.next_steps_actions or data.success_tracking
                   or data.open_questions)
    if not has_content:
        return

    doc.add_page_break()
    _add_heading(doc, "Next Steps", level=2)

    if data.next_steps_actions:
        _add_heading(doc, "Recommended Actions", level=3)
        for line in data.next_steps_actions.strip().split("\n"):
            line = line.strip()
            if line:
                _add_body_paragraph(doc, line)

    if data.success_tracking:
        _add_heading(doc, "Success Tracking", level=3)
        st = data.success_tracking
        _add_body_paragraph(doc, f"Metric: {st.metric}", bold=True)
        _add_body_paragraph(doc, f"Baseline: {st.baseline}")
        _add_body_paragraph(doc, f"Target: {st.target}")
        if st.check_in_date:
            _add_body_paragraph(doc, f"Check-in: {st.check_in_date}")

    if data.open_questions:
        _add_heading(doc, "Open Questions", level=3)
        for line in data.open_questions.strip().split("\n"):
            line = line.strip()
            if line:
                _add_body_paragraph(doc, line)


def _build_resources(doc: Document, data: AnalysisData) -> None:
    """H2 Resources with SQL queries, companion analyses, data sources.

    SQL that is already shown inline with findings is filtered out to avoid
    duplication. If some SQL is inline, the heading becomes "Additional Queries".
    """
    # Determine which SQL is already inline with findings
    inline_sqls: set[str] = set()
    for f in data.findings:
        if f.sql:
            inline_sqls.add(f.sql.strip())

    # Filter to orphan queries (not already inline)
    if inline_sqls:
        orphan_queries = [
            q for q in data.sql_queries
            if q.sql.strip() not in inline_sqls
        ]
    else:
        orphan_queries = list(data.sql_queries)

    has_content = (orphan_queries or data.companion_analyses
                   or data.data_sources)
    if not has_content:
        return

    doc.add_page_break()
    _add_heading(doc, "Resources", level=2)

    # SQL Queries (only orphans)
    if orphan_queries:
        heading = "Additional Queries" if inline_sqls else "Queries"
        _add_heading(doc, heading, level=3)

        for j, query in enumerate(orphan_queries, 1):
            bookmark_name = f"query-{j}"
            _add_heading(doc, f"Query {j}: {query.title}", level=4,
                         bookmark_name=bookmark_name)

            if query.used_in_finding:
                _add_body_paragraph(
                    doc,
                    f"Used in: Finding {query.used_in_finding}",
                    bold=True,
                )

            if query.database:
                _add_body_paragraph(doc, f"Database: {query.database}")

            # SQL in gray-shaded monospace table cell
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.cell(0, 0)
            _set_cell_shading(cell, _SQL_BG_COLOR)

            # Clear default paragraph and add SQL
            cell.text = ""
            p_sql = cell.paragraphs[0]
            run_sql = p_sql.add_run(query.sql.strip())
            _style_run(run_sql, font_name=_SQL_FONT, size=_SQL_SIZE)

    # Companion Analyses
    if data.companion_analyses:
        _add_heading(doc, "Companion Analyses", level=3)
        for line in data.companion_analyses.strip().split("\n"):
            line = line.strip()
            if line:
                p = doc.add_paragraph(line, style="List Bullet")

    # Data Sources
    if data.data_sources:
        _add_heading(doc, "Data Sources", level=3)
        _add_body_paragraph(doc, data.data_sources)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_readout(data: AnalysisData, output_dir: str = "outputs") -> str:
    """Build a .docx Analysis Readout from structured data.

    Args:
        data: All analysis content and metadata.
        output_dir: Directory to write the .docx file.

    Returns:
        Absolute path to the generated .docx file.

    Raises:
        BuildError: If the document cannot be saved (I/O error, permissions).
        Individual section failures are logged and skipped — the document
        will be incomplete but still usable.
    """
    _setup_debug_log()
    _log.info("Starting build_readout: title=%r", data.title)

    doc = Document()

    # Set default font for Normal style
    style = doc.styles["Normal"]
    style.font.name = _BODY_FONT
    style.font.size = _BODY_SIZE
    style.font.color.rgb = _BODY_COLOR

    # Figure counter (mutable list for pass-by-reference)
    figure_counter = [0]
    errors = []

    # Build sections — each wrapped so one failure doesn't kill the doc
    for section_name, builder, args in [
        ("title_block", _build_title_block, (doc, data)),
        ("context", _build_context, (doc, data)),
        ("summary", _build_summary, (doc, data, figure_counter)),
        ("analysis", _build_analysis, (doc, data, figure_counter)),
        ("next_steps", _build_next_steps, (doc, data)),
        ("resources", _build_resources, (doc, data)),
    ]:
        try:
            builder(*args)
        except Exception as e:
            _log.error("Section '%s' failed: %s\n%s",
                       section_name, e, traceback.format_exc())
            errors.append(section_name)
            # Add a visible error marker in the doc
            _add_body_paragraph(
                doc,
                f"[Section '{section_name}' could not be generated: {e}]",
                italic=True,
            )

    if errors:
        _log.warning("Document built with %d section errors: %s",
                     len(errors), errors)

    # Determine output path
    os.makedirs(output_dir, exist_ok=True)
    slug = data.title.lower().replace(" ", "_")[:40]
    # Remove non-alphanumeric chars except underscores
    slug = "".join(c for c in slug if c.isalnum() or c == "_")
    date_str = data.date.replace("-", "").replace("/", "")[:8] if data.date else "undated"
    filename = f"report_{slug}_{date_str}.docx"
    output_path = os.path.join(output_dir, filename)

    try:
        doc.save(output_path)
    except Exception as e:
        _log.error("Failed to save .docx to %s: %s\n%s",
                   output_path, e, traceback.format_exc())
        raise RuntimeError(
            f"Could not save document to {output_path}: {e}"
        ) from e

    _log.info("Document saved: %s (%d bytes, %d section errors)",
              output_path, os.path.getsize(output_path), len(errors))
    return os.path.abspath(output_path)
